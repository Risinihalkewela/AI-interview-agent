"""
╔══════════════════════════════════════════════════════════════════════════════╗
║   AI INTERVIEW COACH + CV OPTIMIZATION AGENT                               ║
║   Single-file Streamlit Application · Powered by Groq API                  ║
║                                                                              ║
║   ARCHITECTURE: True Autonomous AI Agent                                    ║
║   ─────────────────────────────────────────────────────────────────────     ║
║   WHY THIS IS A TRUE AI AGENT (not just a chatbot):                         ║
║                                                                              ║
║   1. PERCEIVE  — Reads job description, CV text, user answers, past         ║
║                  performance data. Builds a structured world model.          ║
║                                                                              ║
║   2. REASON    — Compares CV vs job, identifies skill gaps, evaluates       ║
║                  answer quality, and decides the NEXT BEST ACTION:           ║
║                  what to ask, at what difficulty, on which topic.            ║
║                                                                              ║
║   3. ACT       — Generates tailored questions, evaluates answers,           ║
║                  rewrites the CV, and produces reports.                      ║
║                                                                              ║
║   4. LEARN     — Updates a persistent memory profile after every            ║
║                  interaction. Future behavior changes based on what          ║
║                  the agent has learned about this user.                      ║
║                                                                              ║
║   The agent has GOALS (maximize readiness), MEMORY (persistent JSON),       ║
║   AUTONOMY (no human tells it what question to ask next), and adapts        ║
║   its own behavior over time — the four pillars of an AI agent.             ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

# ══════════════════════════════════════════════════════════════════════════════
# IMPORTS
# ══════════════════════════════════════════════════════════════════════════════
import os, json, re, uuid, datetime, io, copy, textwrap
from pathlib import Path
from typing import Optional

import streamlit as st
import plotly.graph_objects as go
import plotly.express as px
import pandas as pd
from dotenv import load_dotenv

# Groq
from groq import Groq

# CV parsing
import pdfplumber
from docx import Document as DocxDocument

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, KeepTogether
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

# ══════════════════════════════════════════════════════════════════════════════
# ENVIRONMENT SETUP
# ══════════════════════════════════════════════════════════════════════════════
load_dotenv()
MEMORY_FILE  = Path("agent_memory.json")
CV_CACHE_FILE = Path("cv_cache.json")

# Groq model to use — fast and highly capable
GROQ_MODEL = "llama-3.3-70b-versatile"


def get_api_key() -> str:
    """Return the active Groq API key — session state is the ONE source of truth.
    On first run, seeds session state from the .env file if available."""
    # Seed from .env on very first call (before session state is set)
    if "_groq_api_key" not in st.session_state:
        env_key = os.getenv("GROQ_API_KEY", "").strip()
        st.session_state["_groq_api_key"] = env_key  # may be "" — that's fine

    return st.session_state.get("_groq_api_key", "").strip()


def get_groq_client() -> Groq:
    """Initialise Groq client from the current live key."""
    key = get_api_key()
    if not key:
        _render_api_key_screen()
        st.stop()
    return Groq(api_key=key)


def _render_api_key_screen():
    """Full-page API key entry — tests the key live before accepting it."""
    st.markdown("""
    <style>
    section[data-testid="stSidebar"] { display: none !important; }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style="max-width:500px;margin:60px auto 0;padding:48px 40px;
         background:rgba(255,255,255,0.03);border:1px solid rgba(0,255,170,0.25);
         border-radius:24px;text-align:center">
      <div style="font-size:52px;margin-bottom:16px">🔑</div>
      <div style="font-size:24px;font-weight:800;color:#f0f6ff;margin-bottom:10px">
        Groq API Key Required
      </div>
      <div style="color:rgba(200,214,229,0.5);font-size:13px;line-height:1.7">
        This app uses Groq's <strong style="color:rgba(200,214,229,0.8)">free</strong>
        LLaMA 3.3 70B model.<br>
        Get your free key at
        <a href="https://console.groq.com/keys" target="_blank"
           style="color:#00ffaa;text-decoration:none;font-weight:600">
          console.groq.com/keys
        </a>
      </div>
    </div>
    """, unsafe_allow_html=True)

    col_l, col_m, col_r = st.columns([1, 2, 1])
    with col_m:
        st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)

        raw_key = st.text_input(
            "Paste your Groq API key",
            type="password",
            placeholder="gsk_...",
            key="_key_input_widget",
        )

        if st.button("🚀 Connect & Verify", type="primary",
                     use_container_width=True, key="_connect_btn"):
            val = (raw_key or "").strip()
            if not val:
                st.error("⚠️ Please paste your Groq API key above.")
            else:
                with st.spinner("Verifying key with Groq…"):
                    try:
                        test_client = Groq(api_key=val)
                        test_client.chat.completions.create(
                            model=GROQ_MODEL,
                            messages=[{"role": "user", "content": "hi"}],
                            max_tokens=5,
                        )
                        st.session_state["_groq_api_key"] = val
                        st.session_state["groq_client"] = None
                        st.session_state["brain"] = None
                        st.success("✅ Key verified! Loading app…")
                        st.rerun()
                    except Exception as e:
                        err = str(e)
                        if "401" in err or "invalid_api_key" in err or "Invalid API Key" in err:
                            st.error("❌ Invalid API key — please check and try again.")
                        elif "429" in err:
                            st.warning("⚠️ Rate limit hit. Wait a moment then click Connect again.")
                        else:
                            st.error(f"❌ Connection error: {err}")

        st.markdown("""
        <div style="margin-top:16px;padding:12px 16px;background:rgba(0,255,170,0.04);
             border:1px solid rgba(0,255,170,0.15);border-radius:10px;
             font-size:12px;color:rgba(200,214,229,0.5);text-align:left;line-height:1.8">
          <strong style="color:rgba(200,214,229,0.7)">Steps to get your free key:</strong><br>
          1. Go to <a href="https://console.groq.com/keys" target="_blank"
             style="color:#00ffaa;text-decoration:none">console.groq.com/keys</a><br>
          2. Sign up / log in (no credit card needed)<br>
          3. Click <em>Create API Key</em><br>
          4. Copy the key starting with <code style="color:#00ffaa">gsk_</code> and paste above
        </div>
        """, unsafe_allow_html=True)


def groq_chat(client: Groq, prompt: str, system: str = "", max_tokens: int = 4096) -> str:
    """Single helper that wraps every Groq API call uniformly.
    Always rebuilds the client from the current live API key so a stale
    cached client can never cause a 401 error."""
    # Always use the freshest key available — ignores the passed-in client
    live_key = get_api_key()
    if not live_key:
        raise RuntimeError("No Groq API key set. Please enter your key.")
    live_client = Groq(api_key=live_key)

    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    response = live_client.chat.completions.create(
        model=GROQ_MODEL,
        messages=messages,
        max_tokens=max_tokens,
        temperature=0.7,
    )
    return response.choices[0].message.content.strip()


def safe_json(text: str) -> dict:
    """Strip markdown fences and parse JSON robustly."""
    text = re.sub(r"```json|```", "", text).strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        text = match.group(0)
    return json.loads(text)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 · AGENT MEMORY
# ══════════════════════════════════════════════════════════════════════════════
class AgentMemory:
    """
    Persistent user profile — the LEARN component.

    Stores everything the agent knows about the user across ALL sessions.
    This is what differentiates a true agent from a stateless chatbot.

    Schema
    ──────
    user_profile        : name, sessions completed, total questions answered
    weak_skills         : { skill_tag: cumulative_deficit }   ← updated per answer
    scores              : [ {session_id, timestamp, tech, comm, conf, overall} ]
    cv_gaps             : [ skill_string ]                    ← from last CV analysis
    question_history    : [ question_hash ]                   ← deduplication set
    sessions            : [ full session archive ]
    current_session     : live session (questions, answers, evals)
    """

    def __init__(self):
        self._data = self._load()

    # ── Persistence ──────────────────────────────────────────────────────────
    def _load(self) -> dict:
        if MEMORY_FILE.exists():
            try:
                return json.loads(MEMORY_FILE.read_text(encoding="utf-8"))
            except Exception:
                pass
        return self._blank()

    def _blank(self) -> dict:
        return {
            "user_profile": {
                "sessions_completed": 0,
                "total_questions_answered": 0,
                "created_at": datetime.datetime.now().isoformat(),
            },
            "weak_skills": {},
            "strong_skills": {},
            "scores": [],
            "cv_gaps": [],
            "cv_strengths": [],
            "question_history": [],
            "sessions": [],
            "current_session": None,
            "chat_history": [],
        }

    def save(self):
        MEMORY_FILE.write_text(
            json.dumps(self._data, indent=2, default=str),
            encoding="utf-8"
        )

    def reset(self):
        MEMORY_FILE.unlink(missing_ok=True)
        self._data = self._blank()

    # ── Current Session ──────────────────────────────────────────────────────
    def start_session(self, company: str, role: str, level: str,
                      responsibilities: str, requirements: str) -> str:
        sid = str(uuid.uuid4())[:8]
        self._data["current_session"] = {
            "session_id": sid,
            "timestamp": datetime.datetime.now().isoformat(),
            "company": company,
            "role": role,
            "level": level,
            "responsibilities": responsibilities,
            "requirements": requirements,
            "questions": [],
            "answers": [],
            "evaluations": [],
            "scores": [],
            "difficulty": "medium",
            "current_index": 0,
            "company_description": "",
            "question_pool": [],
        }
        self.save()
        return sid

    def close_session(self):
        cs = self._data.get("current_session")
        if not cs:
            return
        if cs.get("scores"):
            sm = self._summarise_scores(cs["scores"])
            cs["summary"] = sm
            self._data["scores"].append({
                "session_id": cs["session_id"],
                "timestamp": cs["timestamp"],
                "company": cs.get("company", ""),
                "role": cs.get("role", ""),
                **sm,
            })
            self._data["user_profile"]["sessions_completed"] += 1
            self._data["user_profile"]["total_questions_answered"] += len(cs["answers"])
        self._data["sessions"].append(copy.deepcopy(cs))
        self._data["current_session"] = None
        self.save()

    @staticmethod
    def _summarise_scores(score_list: list) -> dict:
        if not score_list:
            return {}
        keys = ["technical_knowledge", "communication", "confidence"]
        avgs = {k: round(sum(s.get(k, 0) for s in score_list) / len(score_list), 2)
                for k in keys}
        avgs["overall"] = round(sum(avgs.values()) / 3, 2)
        return avgs

    # ── Learning ──────────────────────────────────────────────────────────────
    def learn_from_answer(self, question: dict, scores: dict):
        tags = question.get("tags", [])
        avg_score = (
            scores.get("technical_knowledge", 3) +
            scores.get("communication", 3) +
            scores.get("confidence", 3)
        ) / 3

        for tag in tags:
            if avg_score < 2.5:
                deficit = 3.0 - avg_score
                self._data["weak_skills"][tag] = (
                    self._data["weak_skills"].get(tag, 0.0) + deficit
                )
            elif avg_score >= 4.0:
                self._data["strong_skills"][tag] = (
                    self._data["strong_skills"].get(tag, 0.0) + (avg_score - 3.0)
                )
                if tag in self._data["weak_skills"]:
                    self._data["weak_skills"][tag] = max(
                        0.0, self._data["weak_skills"][tag] - 0.5
                    )
        self.save()

    def store_cv_analysis(self, gaps: list, strengths: list):
        self._data["cv_gaps"] = gaps
        self._data["cv_strengths"] = strengths
        for gap in gaps:
            tag = gap.lower().replace(" ", "_")
            self._data["weak_skills"][tag] = (
                self._data["weak_skills"].get(tag, 0.0) + 1.5
            )
        self.save()

    # ── Question deduplication ────────────────────────────────────────────────
    def has_asked(self, q_text: str) -> bool:
        h = str(hash(q_text.lower().strip()))
        return h in self._data["question_history"]

    def mark_asked(self, q_text: str):
        h = str(hash(q_text.lower().strip()))
        if h not in self._data["question_history"]:
            self._data["question_history"].append(h)
        self.save()

    # ── Helpers ───────────────────────────────────────────────────────────────
    def update_session(self, key: str, value):
        if self._data["current_session"] is not None:
            self._data["current_session"][key] = value
            self.save()

    def record_qa(self, question: dict, answer: str, evaluation: dict, scores: dict):
        cs = self._data["current_session"]
        if cs:
            cs["questions"].append(question)
            cs["answers"].append(answer)
            cs["evaluations"].append(evaluation)
            cs["scores"].append(scores)
            cs["current_index"] += 1
        self.save()

    def add_chat(self, role: str, content: str):
        self._data["chat_history"].append({"role": role, "content": content})
        self._data["chat_history"] = self._data["chat_history"][-60:]
        self.save()

    def clear_chat(self):
        self._data["chat_history"] = []
        self.save()

    # ── Accessors ─────────────────────────────────────────────────────────────
    @property
    def weak_skills(self) -> dict:
        return self._data.get("weak_skills", {})

    @property
    def strong_skills(self) -> dict:
        return self._data.get("strong_skills", {})

    @property
    def cv_gaps(self) -> list:
        return self._data.get("cv_gaps", [])

    @property
    def cv_strengths(self) -> list:
        return self._data.get("cv_strengths", [])

    @property
    def sessions(self) -> list:
        return self._data.get("sessions", [])

    @property
    def current_session(self) -> Optional[dict]:
        return self._data.get("current_session")

    @property
    def chat_history(self) -> list:
        return self._data.get("chat_history", [])

    @property
    def score_history(self) -> list:
        return self._data.get("scores", [])

    @property
    def user_profile(self) -> dict:
        return self._data.get("user_profile", {})


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 · AGENT BRAIN (PERCEIVE → REASON → ACT → LEARN)
# ══════════════════════════════════════════════════════════════════════════════
class AgentBrain:
    """
    The autonomous decision-making core of the agent.

    DECISION LOGIC:
    ┌──────────────────────────────────────────────────────────────┐
    │ IF technical_score < 3  → prioritise technical questions     │
    │ IF communication < 3    → add HR / behavioural questions     │
    │ IF CV has gap X         → generate questions about X         │
    │ IF avg_score ≥ 4.0      → increase difficulty                │
    │ IF avg_score ≤ 2.0      → decrease difficulty                │
    │ IF question seen before → skip, pick alternative             │
    │ IF weak_skill high      → target that skill first            │
    └──────────────────────────────────────────────────────────────┘
    """

    def __init__(self, client: Groq, memory: AgentMemory):
        self.client = client
        self.memory = memory

    # ── PERCEIVE ──────────────────────────────────────────────────────────────
    def perceive(self) -> dict:
        cs = self.memory.current_session or {}
        scores_so_far = cs.get("scores", [])

        if scores_so_far:
            avg_tech = sum(s.get("technical_knowledge", 3) for s in scores_so_far) / len(scores_so_far)
            avg_comm = sum(s.get("communication", 3) for s in scores_so_far) / len(scores_so_far)
            avg_conf = sum(s.get("confidence", 3) for s in scores_so_far) / len(scores_so_far)
        else:
            avg_tech = avg_comm = avg_conf = 3.0

        top_weak = sorted(
            self.memory.weak_skills.items(), key=lambda x: x[1], reverse=True
        )[:5]

        return {
            "role": cs.get("role", ""),
            "company": cs.get("company", ""),
            "level": cs.get("level", ""),
            "responsibilities": cs.get("responsibilities", ""),
            "requirements": cs.get("requirements", ""),
            "current_difficulty": cs.get("difficulty", "medium"),
            "questions_answered": len(scores_so_far),
            "avg_technical": round(avg_tech, 2),
            "avg_communication": round(avg_comm, 2),
            "avg_confidence": round(avg_conf, 2),
            "top_weak_skills": top_weak,
            "cv_gaps": self.memory.cv_gaps,
            "cv_strengths": self.memory.cv_strengths,
            "session_count": self.memory.user_profile.get("sessions_completed", 0),
        }

    # ── REASON ────────────────────────────────────────────────────────────────
    def reason(self, perception: dict) -> dict:
        avg_t = perception["avg_technical"]
        avg_c = perception["avg_communication"]
        avg_conf = perception["avg_confidence"]
        n_answered = perception["questions_answered"]
        current_diff = perception["current_difficulty"]
        top_weak = perception["top_weak_skills"]
        cv_gaps = perception["cv_gaps"]

        decision = {
            "next_action": "generate_questions",
            "question_type": "mixed",
            "target_skills": [],
            "difficulty": current_diff,
            "rationale": [],
        }

        if avg_t < 3.0 and n_answered >= 3:
            decision["question_type"] = "technical"
            decision["rationale"].append(
                f"Technical avg {avg_t:.1f} < 3.0 → prioritising technical questions"
            )
        elif avg_c < 3.0 and n_answered >= 3:
            decision["question_type"] = "hr"
            decision["rationale"].append(
                f"Communication avg {avg_c:.1f} < 3.0 → adding HR/behavioural questions"
            )

        if cv_gaps:
            skill_tags = [g.lower().replace(" ", "_") for g in cv_gaps[:3]]
            decision["target_skills"].extend(skill_tags)
            decision["rationale"].append(
                f"CV gaps detected: {', '.join(cv_gaps[:3])} → targeting these skills"
            )

        if top_weak:
            for tag, deficit in top_weak[:2]:
                if tag not in decision["target_skills"]:
                    decision["target_skills"].append(tag)
            decision["rationale"].append(
                f"Weak areas: {[t for t,_ in top_weak[:2]]} → questioning these topics"
            )

        overall_avg = (avg_t + avg_c + avg_conf) / 3
        if n_answered >= 5:
            if overall_avg >= 4.2 and current_diff != "hard":
                decision["difficulty"] = "hard"
                decision["next_action"] = "adapt_difficulty"
                decision["rationale"].append(
                    f"Overall avg {overall_avg:.1f} ≥ 4.2 → escalating to HARD"
                )
            elif overall_avg <= 2.2 and current_diff != "easy":
                decision["difficulty"] = "easy"
                decision["next_action"] = "adapt_difficulty"
                decision["rationale"].append(
                    f"Overall avg {overall_avg:.1f} ≤ 2.2 → de-escalating to EASY"
                )
            elif 2.2 < overall_avg < 4.2 and current_diff == "hard":
                decision["difficulty"] = "medium"
                decision["rationale"].append("Score recovering → resetting to MEDIUM")

        if not decision["rationale"]:
            decision["rationale"].append("Balanced performance → standard progression")

        return decision

    # ── ACT: Research Company ─────────────────────────────────────────────────
    def act_research_company(self, company: str, role: str) -> str:
        prompt = (
            f'Provide a 3-4 sentence overview of "{company}" covering: what they do, '
            f'their culture/values, and why someone would want to work there in the role '
            f'of "{role}". Be factual. If unknown, write a plausible generic profile.'
        )
        try:
            return groq_chat(
                self.client, prompt,
                system="You are a professional researcher.",
                max_tokens=300
            )
        except Exception as e:
            return f"A leading organisation in their industry. (Research unavailable: {e})"

    # ── ACT: Generate Questions ───────────────────────────────────────────────
    def act_generate_questions(
        self, perception: dict, decision: dict, count: int = 20
    ) -> dict:
        n_hr   = max(4, count // 4)
        n_tech = count - n_hr

        weak_str = ", ".join(
            f"{t} (deficit:{v:.1f})" for t, v in perception["top_weak_skills"]
        ) or "None yet"

        cv_gap_str = ", ".join(perception["cv_gaps"][:5]) or "None provided"
        target_str = ", ".join(decision["target_skills"]) or "general topics"

        q_type_instruction = {
            "technical": f"Skew heavily toward technical questions. Generate {n_hr} HR and {n_tech} technical.",
            "hr":        f"Include extra HR/behavioural. Generate {n_hr + 3} HR and {n_tech - 3} technical.",
            "mixed":     f"Generate {n_hr} HR and {n_tech} technical.",
        }.get(decision["question_type"], f"Generate {n_hr} HR and {n_tech} technical.")

        diff_guide = {
            "easy":   "straightforward, foundational questions suitable for beginners",
            "medium": "scenario-based, moderately complex questions",
            "hard":   "deep multi-part questions, edge cases, system design, advanced concepts",
        }.get(decision["difficulty"], "scenario-based questions")

        rationale_str = " | ".join(decision.get("rationale", []))

        prompt = f"""You are a senior technical interviewer. Generate exactly {count} interview questions.
Return ONLY valid JSON with no markdown or preamble.

CANDIDATE PROFILE:
- Company: {perception['company']}
- Role: {perception['role']}
- Level: {perception['level']}
- Responsibilities: {perception['responsibilities']}
- Requirements: {perception['requirements']}
- Difficulty: {decision['difficulty']} ({diff_guide})
- Agent's Target Skills: {target_str}
- Candidate's Weak Skills: {weak_str}
- CV Gaps to Address: {cv_gap_str}

INSTRUCTIONS:
{q_type_instruction}
- Questions MUST be specific to the role and company
- Weak skills and CV gaps should appear in multiple questions
- Each question needs skill tags for memory tracking
- Include a short "follow_up_hint" the interviewer would ask

Return EXACTLY this JSON (no other text):
{{
  "hr_questions": [
    {{
      "id": "hr_1",
      "question": "...",
      "type": "hr",
      "difficulty": "{decision['difficulty']}",
      "tags": ["communication", "teamwork"],
      "follow_up_hint": "..."
    }}
  ],
  "technical_questions": [
    {{
      "id": "tech_1",
      "question": "...",
      "type": "technical",
      "difficulty": "{decision['difficulty']}",
      "tags": ["python", "machine_learning"],
      "follow_up_hint": "..."
    }}
  ],
  "generation_rationale": "{rationale_str}"
}}"""
        try:
            raw = groq_chat(self.client, prompt, max_tokens=4096)
            return safe_json(raw)
        except Exception as e:
            st.error(f"Question generation failed: {e}")
            return {"hr_questions": [], "technical_questions": [], "generation_rationale": ""}

    # ── ACT: Evaluate Answer ──────────────────────────────────────────────────
    def act_evaluate_answer(
        self, question: dict, answer: str, role: str, level: str, context: str = ""
    ) -> dict:
        prompt = f"""You are a senior interviewer evaluating a candidate for {role} at {level} level.
Return ONLY valid JSON with no markdown or preamble.

QUESTION: {question.get('question', '')}
TYPE: {question.get('type', 'general')}
DIFFICULTY: {question.get('difficulty', 'medium')}
SKILL TAGS: {', '.join(question.get('tags', []))}
CANDIDATE ANSWER: {answer}
CONTEXT: {context}

Evaluate rigorously. A non-answer or "I don't know" must score 1.

Return EXACTLY this JSON:
{{
  "strengths": ["...", "..."],
  "weaknesses": ["...", "..."],
  "suggestions": ["...", "...", "..."],
  "ideal_answer_summary": "2-3 sentences on what a perfect answer covers",
  "scores": {{
    "technical_knowledge": <1-5>,
    "communication": <1-5>,
    "confidence": <1-5>
  }},
  "overall_comment": "One paragraph assessment",
  "skill_demonstrated": true
}}

SCORING RUBRIC:
1 = Very Poor (barely relevant or blank)
2 = Poor (missing key concepts)
3 = Average (correct but shallow)
4 = Good (solid, well-structured)
5 = Excellent (comprehensive, impressive depth)"""
        try:
            raw = groq_chat(self.client, prompt, max_tokens=1500)
            data = safe_json(raw)
            sc = data.get("scores", {})
            for k in ["technical_knowledge", "communication", "confidence"]:
                sc[k] = max(1, min(5, int(sc.get(k, 3))))
            data["scores"] = sc
            return data
        except Exception as e:
            return {
                "strengths": ["Could not parse evaluation"],
                "weaknesses": ["Please try again"],
                "suggestions": ["Check API connection"],
                "ideal_answer_summary": "",
                "scores": {"technical_knowledge": 3, "communication": 3, "confidence": 3},
                "overall_comment": f"Evaluation error: {e}",
                "skill_demonstrated": False,
            }

    # ── ACT: Analyse CV ───────────────────────────────────────────────────────
    def act_analyse_cv(
        self, cv_text: str, role: str, level: str,
        responsibilities: str, requirements: str
    ) -> dict:
        prompt = f"""You are an expert ATS-optimised CV consultant and career coach.
Analyse the CV against the job requirements and return ONLY valid JSON.

JOB DETAILS:
- Role: {role}
- Level: {level}
- Responsibilities: {responsibilities}
- Requirements: {requirements}

CV CONTENT:
{cv_text[:4000]}

Return EXACTLY this JSON (no markdown):
{{
  "relevance_score": <1-10>,
  "strengths": ["...", "..."],
  "weaknesses": ["...", "..."],
  "missing_skills": ["skill1", "skill2", "skill3"],
  "ats_issues": ["issue1", "issue2"],
  "keywords_missing": ["keyword1", "keyword2"],
  "keywords_present": ["keyword1", "keyword2"],
  "experience_gaps": ["gap1", "gap2"],
  "recommendations": ["rec1", "rec2", "rec3"],
  "overall_assessment": "2-3 sentences",
  "match_breakdown": {{
    "technical_match": <1-10>,
    "experience_match": <1-10>,
    "keyword_match": <1-10>
  }}
}}"""
        try:
            raw = groq_chat(self.client, prompt, max_tokens=2000)
            return safe_json(raw)
        except Exception as e:
            return {"relevance_score": 0, "error": str(e)}

    # ── ACT: Rewrite CV ───────────────────────────────────────────────────────
    def act_rewrite_cv(
        self, cv_text: str, role: str, level: str,
        responsibilities: str, requirements: str,
        analysis: dict
    ) -> str:
        missing  = ", ".join(analysis.get("missing_skills", []))
        keywords = ", ".join(analysis.get("keywords_missing", []))
        weaknesses = ", ".join(analysis.get("weaknesses", []))

        prompt = f"""You are an expert CV writer specialising in ATS optimisation.

TASK: Rewrite the CV below to be perfectly aligned with the job of {role} at {level} level.

RULES (STRICT):
1. Do NOT invent experience, companies, degrees, or skills the person does not have
2. DO improve bullet points to use strong action verbs and quantify achievements where possible
3. DO add a tailored Professional Summary at the top
4. DO naturally incorporate missing keywords: {keywords}
5. DO restructure sections for maximum impact
6. DO address these CV weaknesses: {weaknesses}

JOB REQUIREMENTS: {requirements}
KEY RESPONSIBILITIES: {responsibilities}

ORIGINAL CV:
{cv_text[:4000]}

Output ONLY the rewritten CV text in clean plain text format.
Use this structure:
[FULL NAME]
[Contact: Email | Phone | LinkedIn | Location]

PROFESSIONAL SUMMARY
[3-4 compelling sentences]

CORE SKILLS
[Comma-separated or bullet list of key skills]

WORK EXPERIENCE
[Company | Role | Dates]
• [Achievement bullet]
• [Achievement bullet]

EDUCATION
[Degree | Institution | Year]

CERTIFICATIONS (if any)
[Cert | Issuer | Year]"""
        try:
            return groq_chat(self.client, prompt, max_tokens=3000)
        except Exception as e:
            return f"CV rewrite failed: {e}"

    # ── ACT: Chatbot ──────────────────────────────────────────────────────────
    def act_chat(self, user_msg: str, history: list, context: str = "") -> str:
        system = (
            "You are an expert AI Interview Coach. Answer questions about "
            "interview techniques, career strategy, salary negotiation, resume tips, and "
            "professional development. Be concise, specific, and encouraging."
            + (f"\nCurrent session: {context}" if context else "")
        )
        messages = [{"role": "system", "content": system}]
        for m in history[-8:]:
            messages.append({"role": m["role"], "content": m["content"]})
        messages.append({"role": "user", "content": user_msg})

        try:
            live_key = get_api_key()
            if not live_key:
                return "Error: No API key set. Please enter your Groq API key."
            live_client = Groq(api_key=live_key)
            response = live_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=messages,
                max_tokens=800,
                temperature=0.7,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            return f"Error: {e}"

    # ── LEARN ─────────────────────────────────────────────────────────────────
    def learn(self, question: dict, scores: dict):
        self.memory.learn_from_answer(question, scores)
        self._adapt_difficulty_midstream(scores)

    def _adapt_difficulty_midstream(self, latest_scores: dict):
        cs = self.memory.current_session
        if not cs:
            return
        all_scores = cs.get("scores", [])
        if len(all_scores) % 5 != 0 or len(all_scores) == 0:
            return
        recent = all_scores[-5:]
        avgs = [(s["technical_knowledge"] + s["communication"] + s["confidence"]) / 3
                for s in recent]
        session_avg = sum(avgs) / len(avgs)
        current = cs.get("difficulty", "medium")
        new_diff = current
        if session_avg >= 4.2 and current != "hard":
            new_diff = "hard"
            st.toast("🔥 You're excelling — difficulty raised to **Hard**!", icon="📈")
        elif session_avg <= 2.0 and current != "easy":
            new_diff = "easy"
            st.toast("💡 Difficulty adjusted to **Easy** to build confidence.", icon="📉")
        if new_diff != current:
            self.memory.update_session("difficulty", new_diff)

    # ── Select Next Question ──────────────────────────────────────────────────
    def select_next_question(self, pool: list) -> Optional[dict]:
        top_weak_tags = set(t for t, _ in
                            sorted(self.memory.weak_skills.items(),
                                   key=lambda x: x[1], reverse=True)[:5])
        cv_gap_tags = set(g.lower().replace(" ", "_") for g in self.memory.cv_gaps)

        unseen = [q for q in pool if not self.memory.has_asked(q.get("question", ""))]
        if not unseen:
            unseen = pool

        def score_q(q):
            tags = set(q.get("tags", []))
            return len(tags & top_weak_tags) * 3 + len(tags & cv_gap_tags) * 2

        ranked = sorted(unseen, key=score_q, reverse=True)
        return ranked[0] if ranked else None


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 · CV PARSING
# ══════════════════════════════════════════════════════════════════════════════
def extract_cv_text(uploaded_file) -> str:
    filename  = uploaded_file.name.lower()
    raw_bytes = uploaded_file.read()

    if filename.endswith(".pdf"):
        try:
            with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                pages = [p.extract_text() or "" for p in pdf.pages]
            text = "\n".join(pages).strip()
            if text:
                return text
        except Exception as e:
            st.warning(f"PDF parse warning: {e}")
        return ""

    elif filename.endswith(".docx"):
        try:
            doc = DocxDocument(io.BytesIO(raw_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            st.warning(f"DOCX parse warning: {e}")
        return ""

    return "Unsupported file format."


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 · PDF REPORT GENERATOR
# ══════════════════════════════════════════════════════════════════════════════
def generate_pdf_report(memory: AgentMemory, session_data: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    S = getSampleStyleSheet()
    BLUE  = colors.HexColor("#1e40af")
    NAVY  = colors.HexColor("#0f172a")
    GREY  = colors.HexColor("#64748b")
    LGREY = colors.HexColor("#f8fafc")
    GREEN = colors.HexColor("#15803d")
    RED   = colors.HexColor("#dc2626")
    AMBER = colors.HexColor("#d97706")

    TITLE  = ParagraphStyle("T", fontSize=22, textColor=NAVY, alignment=TA_CENTER,
                            fontName="Helvetica-Bold", spaceAfter=4)
    SUB    = ParagraphStyle("S", fontSize=13, textColor=GREY, alignment=TA_CENTER,
                            spaceAfter=16)
    H1     = ParagraphStyle("H1", fontSize=14, textColor=BLUE,
                            fontName="Helvetica-Bold", spaceBefore=14, spaceAfter=6)
    H2     = ParagraphStyle("H2", fontSize=11, textColor=NAVY,
                            fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=4)
    BODY   = ParagraphStyle("B", fontSize=9.5, leading=14, spaceAfter=4)
    SMALL  = ParagraphStyle("SM", fontSize=8, textColor=GREY, spaceAfter=2)
    BULLET = ParagraphStyle("BU", fontSize=9.5, leading=14, leftIndent=15, spaceAfter=2)

    story = []

    story += [
        Spacer(1, 1.5*cm),
        Paragraph("🎯 AI Interview Coach", TITLE),
        Paragraph("Performance & CV Optimization Report", SUB),
        HRFlowable(width="100%", thickness=2, color=BLUE),
        Spacer(1, 0.4*cm),
    ]

    cs = session_data
    summary = cs.get("summary", {})
    meta = [
        ["Company", cs.get("company","—"), "Role", cs.get("role","—")],
        ["Level",   cs.get("level","—"),   "Date", cs.get("timestamp","")[:10]],
        ["Questions Answered", str(len(cs.get("answers",[]))),
         "Session ID", cs.get("session_id","—")],
    ]
    meta_tbl = Table(meta, colWidths=[3.5*cm, 6*cm, 3.5*cm, 3.5*cm])
    meta_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0),(0,-1), colors.HexColor("#dbeafe")),
        ("BACKGROUND", (2,0),(2,-1), colors.HexColor("#dbeafe")),
        ("FONTNAME",  (0,0),(0,-1), "Helvetica-Bold"),
        ("FONTNAME",  (2,0),(2,-1), "Helvetica-Bold"),
        ("FONTSIZE",  (0,0),(-1,-1), 9),
        ("GRID",      (0,0),(-1,-1), 0.5, colors.HexColor("#cbd5e1")),
        ("PADDING",   (0,0),(-1,-1), 6),
        ("BACKGROUND",(1,0),(-1,-1), LGREY),
    ]))
    story += [meta_tbl, Spacer(1, 0.5*cm)]

    if summary:
        story.append(Paragraph("📊 Overall Performance", H1))
        def _rating(v):
            if v >= 4.5: return ("Excellent", GREEN)
            if v >= 3.5: return ("Good", BLUE)
            if v >= 2.5: return ("Average", AMBER)
            return ("Needs Work", RED)
        rows = [["Dimension","Score","/ 5","Rating"]]
        for dim, label in [
            ("technical_knowledge","Technical Knowledge"),
            ("communication","Communication"),
            ("confidence","Confidence"),
            ("overall","🏆 OVERALL"),
        ]:
            val = summary.get(dim,0)
            rat, col = _rating(val)
            rows.append([label, f"{val:.2f}", "/ 5",
                         Paragraph(f'<font color="{col.hexval()}">{rat}</font>', BODY)])
        sc_tbl = Table(rows, colWidths=[7*cm, 3.5*cm, 2*cm, 4*cm])
        sc_tbl.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0), BLUE),
            ("TEXTCOLOR", (0,0),(-1,0), colors.white),
            ("FONTNAME",  (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTNAME",  (0,4),(-1,4), "Helvetica-Bold"),
            ("BACKGROUND",(0,4),(-1,4), colors.HexColor("#eff6ff")),
            ("GRID",      (0,0),(-1,-1), 0.5, colors.HexColor("#bfdbfe")),
            ("FONTSIZE",  (0,0),(-1,-1), 10),
            ("ALIGN",     (1,0),(-1,-1), "CENTER"),
            ("PADDING",   (0,0),(-1,-1), 7),
        ]))
        story += [sc_tbl, Spacer(1, 0.5*cm)]

    if memory.cv_gaps:
        story.append(Paragraph("📄 CV Gaps Identified", H1))
        for g in memory.cv_gaps:
            story.append(Paragraph(f"• {g}", BULLET))
        story.append(Spacer(1, 0.3*cm))

    story.append(PageBreak())
    story.append(Paragraph("📝 Question-by-Question Breakdown", H1))
    for i, (q, a, ev, sc) in enumerate(zip(
        cs.get("questions",[]), cs.get("answers",[]),
        cs.get("evaluations",[]), cs.get("scores",[])
    )):
        block = [
            Paragraph(f"Q{i+1}. [{q.get('type','').upper()}] {q.get('question','')}", H2),
            Paragraph(f"Tags: {', '.join(q.get('tags',[]))}  |  Difficulty: {q.get('difficulty','')}", SMALL),
            Spacer(1, 0.1*cm),
            Paragraph(f"<b>Your Answer:</b> {a or '—'}", BODY),
            Spacer(1, 0.1*cm),
        ]
        sc_row = Table(
            [["Technical","Communication","Confidence"],
             [f"{sc.get('technical_knowledge',0)}/5",
              f"{sc.get('communication',0)}/5",
              f"{sc.get('confidence',0)}/5"]],
            colWidths=[5.5*cm, 5.5*cm, 5.5*cm]
        )
        sc_row.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0), colors.HexColor("#dbeafe")),
            ("FONTNAME",  (0,0),(-1,0), "Helvetica-Bold"),
            ("ALIGN",     (0,0),(-1,-1), "CENTER"),
            ("FONTSIZE",  (0,0),(-1,-1), 9),
            ("GRID",      (0,0),(-1,-1), 0.5, colors.HexColor("#bfdbfe")),
            ("PADDING",   (0,0),(-1,-1), 5),
        ]))
        block.append(sc_row)
        block.append(Spacer(1, 0.1*cm))
        if ev.get("strengths"):
            block.append(Paragraph("<b>✅ Strengths:</b>", BODY))
            for s in ev["strengths"]: block.append(Paragraph(f"• {s}", BULLET))
        if ev.get("weaknesses"):
            block.append(Paragraph("<b>⚠️ Areas to Improve:</b>", BODY))
            for w in ev["weaknesses"]: block.append(Paragraph(f"• {w}", BULLET))
        if ev.get("suggestions"):
            block.append(Paragraph("<b>💡 Suggestions:</b>", BODY))
            for sg in ev["suggestions"]: block.append(Paragraph(f"• {sg}", BULLET))
        if ev.get("ideal_answer_summary"):
            block.append(Paragraph(f"<b>📌 Ideal Answer:</b> {ev['ideal_answer_summary']}", BODY))
        block.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e2e8f0")))
        block.append(Spacer(1, 0.15*cm))
        story.append(KeepTogether(block))

    story += [
        PageBreak(),
        Paragraph("🎯 Agent Memory: Top Weak Skills", H1),
    ]
    ws = sorted(memory.weak_skills.items(), key=lambda x: x[1], reverse=True)[:10]
    if ws:
        wa_rows = [["Skill","Deficit Score","Priority"]] + [
            [k.replace("_"," ").title(), f"{v:.1f}",
             "🔴 High" if v > 3 else "🟡 Medium" if v > 1 else "🟢 Low"]
            for k,v in ws
        ]
        wa_tbl = Table(wa_rows, colWidths=[8*cm, 4*cm, 4.5*cm])
        wa_tbl.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0), colors.HexColor("#7c3aed")),
            ("TEXTCOLOR", (0,0),(-1,0), colors.white),
            ("FONTNAME",  (0,0),(-1,0), "Helvetica-Bold"),
            ("GRID",      (0,0),(-1,-1), 0.5, colors.HexColor("#ddd6fe")),
            ("FONTSIZE",  (0,0),(-1,-1), 9),
            ("ALIGN",     (1,0),(-1,-1), "CENTER"),
            ("PADDING",   (0,0),(-1,-1), 6),
        ]))
        story.append(wa_tbl)

    story += [
        Spacer(1, 1*cm),
        HRFlowable(width="100%", thickness=1, color=GREY),
        Paragraph(
            f"Generated by AI Interview Coach Agent · {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}",
            ParagraphStyle("ft", fontSize=8, textColor=GREY, alignment=TA_CENTER, spaceBefore=6)
        ),
    ]

    doc.build(story)
    return buf.getvalue()


def generate_cv_pdf(cv_text: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    BLUE = colors.HexColor("#1e40af")
    story = []
    for line in cv_text.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.2*cm))
            continue
        if line.isupper() and len(line) < 50:
            story.append(Paragraph(line, ParagraphStyle(
                "sh", fontSize=11, fontName="Helvetica-Bold",
                textColor=BLUE, spaceBefore=10, spaceAfter=4
            )))
            story.append(HRFlowable(width="100%", thickness=0.5, color=BLUE))
        elif line.startswith("•"):
            story.append(Paragraph(line, ParagraphStyle(
                "bu", fontSize=9.5, leading=14, leftIndent=10, spaceAfter=2
            )))
        else:
            story.append(Paragraph(line, ParagraphStyle(
                "bo", fontSize=9.5, leading=14, spaceAfter=3
            )))
    doc.build(story)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 · VISUALISATION
# ══════════════════════════════════════════════════════════════════════════════
def score_bars(scores: dict):
    config = {
        "technical_knowledge": ("🔧 Technical Knowledge", "#00ffaa"),
        "communication":       ("🗣️ Communication",       "#a5b4fc"),
        "confidence":          ("💪 Confidence",           "#fbbf24"),
    }
    for key, (label, colour) in config.items():
        v = scores.get(key, 0)
        pct = v / 5 * 100
        st.markdown(f"""
        <div style="margin-bottom:16px">
          <div style="display:flex;justify-content:space-between;margin-bottom:6px">
            <span style="font-size:13px;font-weight:500;color:rgba(200,214,229,0.7)">{label}</span>
            <span style="font-size:14px;font-weight:700;color:{colour}">{v}/5</span>
          </div>
          <div style="background:rgba(255,255,255,0.06);border-radius:999px;height:6px;overflow:hidden">
            <div style="width:{pct}%;background:linear-gradient(90deg,{colour},{colour}99);
                 height:100%;border-radius:999px;transition:width 0.8s cubic-bezier(.4,0,.2,1)"></div>
          </div>
        </div>""", unsafe_allow_html=True)


def radar_chart(scores: dict, title: str = "") -> go.Figure:
    cats = ["Technical", "Communication", "Confidence"]
    vals = [scores.get("technical_knowledge", 0),
            scores.get("communication", 0),
            scores.get("confidence", 0)]
    fig = go.Figure(go.Scatterpolar(
        r=vals + [vals[0]], theta=cats + [cats[0]],
        fill="toself",
        fillcolor="rgba(0,255,170,0.1)",
        line=dict(color="#00ffaa", width=2.5),
        marker=dict(size=8, color="#00ffaa"),
    ))
    fig.update_layout(
        polar=dict(
            radialaxis=dict(visible=True, range=[0, 5], gridcolor="rgba(255,255,255,0.08)", color="rgba(200,214,229,0.4)"),
            angularaxis=dict(gridcolor="rgba(255,255,255,0.08)", color="rgba(200,214,229,0.5)"),
            bgcolor="rgba(0,0,0,0)",
        ),
        showlegend=False, title=dict(text=title, x=0.5, font=dict(size=14, color="#f0f6ff")),
        height=320, margin=dict(l=40,r=40,t=50,b=30),
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="rgba(200,214,229,0.7)"),
    )
    return fig


def trend_chart(score_history: list) -> Optional[go.Figure]:
    if len(score_history) < 2:
        return None
    labels = [f"S{i+1}" for i in range(len(score_history))]
    fig = go.Figure()
    for name, key, colour in [
        ("Technical",    "technical_knowledge", "#00ffaa"),
        ("Communication","communication",        "#a5b4fc"),
        ("Confidence",   "confidence",           "#fbbf24"),
        ("Overall",      "overall",              "#f87171"),
    ]:
        fig.add_trace(go.Scatter(
            x=labels,
            y=[s.get(key, 0) for s in score_history],
            mode="lines+markers", name=name,
            line=dict(color=colour, width=3 if name=="Overall" else 2),
            marker=dict(size=9 if name=="Overall" else 6),
        ))
    fig.update_layout(
        title=dict(text="📈 Performance Trend", font=dict(color="#f0f6ff")),
        xaxis=dict(title="Session", gridcolor="rgba(255,255,255,0.06)", color="rgba(200,214,229,0.5)"),
        yaxis=dict(title="Score", range=[0, 5.5], gridcolor="rgba(255,255,255,0.06)", color="rgba(200,214,229,0.5)"),
        legend=dict(orientation="h", y=1.1, x=0, font=dict(color="rgba(200,214,229,0.7)")),
        height=380, hovermode="x unified",
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="rgba(200,214,229,0.7)"),
    )
    return fig


def weakness_chart(weak_skills: dict) -> Optional[go.Figure]:
    if not weak_skills:
        return None
    items = sorted(weak_skills.items(), key=lambda x: x[1], reverse=True)[:12]
    labels = [k.replace("_"," ").title() for k,_ in items]
    values = [v for _,v in items]
    fig = go.Figure(go.Bar(
        x=values, y=labels, orientation="h",
        marker=dict(color=values, colorscale=[[0,"#00ffaa"],[0.5,"#fbbf24"],[1,"#f87171"]]),
        text=[f"{v:.1f}" for v in values], textposition="outside",
        textfont=dict(color="rgba(200,214,229,0.7)"),
    ))
    fig.update_layout(
        title=dict(text="⚠️ Weak Skill Areas", font=dict(color="#f0f6ff")),
        xaxis=dict(title="Deficit Score", gridcolor="rgba(255,255,255,0.06)", color="rgba(200,214,229,0.5)"),
        yaxis=dict(color="rgba(200,214,229,0.7)"),
        height=max(260, len(labels)*36 + 80),
        margin=dict(l=140, r=80, t=50, b=40),
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="rgba(200,214,229,0.7)"),
    )
    return fig


def cv_match_gauge(score: int) -> go.Figure:
    fig = go.Figure(go.Indicator(
        mode="gauge+number+delta",
        value=score,
        title={"text": "CV–Job Match", "font": {"size": 14, "color": "#f0f6ff"}},
        delta={"reference": 5},
        number={"font": {"color": "#00ffaa", "size": 48}},
        gauge={
            "axis": {"range": [0, 10], "tickcolor": "rgba(200,214,229,0.4)"},
            "bar": {"color": "#00ffaa"},
            "bgcolor": "rgba(0,0,0,0)",
            "bordercolor": "rgba(255,255,255,0.1)",
            "steps": [
                {"range": [0, 4],  "color": "rgba(239,68,68,0.15)"},
                {"range": [4, 7],  "color": "rgba(245,158,11,0.15)"},
                {"range": [7, 10], "color": "rgba(0,255,170,0.15)"},
            ],
            "threshold": {"line": {"color": "#f87171", "width": 2}, "value": 5},
        },
    ))
    fig.update_layout(
        height=260, margin=dict(l=20, r=20, t=40, b=20),
        paper_bgcolor="rgba(0,0,0,0)",
        font=dict(color="rgba(200,214,229,0.7)"),
    )
    return fig


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 · PAGE CONFIG + GLOBAL CSS
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="AI Interview Coach Agent",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded",
)

CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;500;600;700;800&family=Inter:wght@300;400;500;600&display=swap');

* { font-family: 'Inter', sans-serif !important; box-sizing: border-box; }
h1,h2,h3,h4,h5,h6 { font-family: 'Syne', sans-serif !important; }

/* ── Hide Streamlit chrome ── */
#MainMenu, footer, header { visibility: hidden; }
::-webkit-scrollbar { width: 5px; }
::-webkit-scrollbar-thumb { background: rgba(0,255,170,0.3); border-radius: 99px; }

/* ── Top shimmer bar ── */
.stApp::before {
    content: '';
    position: fixed; top: 0; left: 0; right: 0; height: 2px; z-index: 9999;
    background: linear-gradient(90deg, #00ffaa, #6366f1, #00ffaa);
    background-size: 200% 100%;
    animation: shimmer 4s linear infinite;
}
@keyframes shimmer { 0% { background-position: 200% 0; } 100% { background-position: -200% 0; } }
@keyframes pulse-dot { 0%, 100% { opacity: 1; } 50% { opacity: 0.4; } }

/* ══ FIX 1: TABS — labels never overlap, wrap gracefully on narrow screens ══ */
[data-testid="stTabs"] [role="tablist"] {
    display: flex !important;
    flex-wrap: wrap !important;
    gap: 2px !important;
    overflow: visible !important;
    border-bottom: 1px solid rgba(255,255,255,0.08) !important;
    padding-bottom: 0 !important;
}
[data-testid="stTabs"] [role="tab"] {
    flex: 0 1 auto !important;
    min-width: 0 !important;
    max-width: 100% !important;
    white-space: nowrap !important;
    overflow: hidden !important;
    text-overflow: ellipsis !important;
    padding: 10px 18px !important;
    font-size: 13px !important;
    font-weight: 600 !important;
    letter-spacing: 0.01em !important;
    border-radius: 8px 8px 0 0 !important;
    transition: background 0.2s ease !important;
}
[data-testid="stTabs"] [role="tab"][aria-selected="true"] {
    background: rgba(0,255,170,0.08) !important;
    border-bottom: 2px solid #00ffaa !important;
    color: #00ffaa !important;
}
[data-testid="stTabs"] [role="tab"]:hover:not([aria-selected="true"]) {
    background: rgba(255,255,255,0.04) !important;
}
/* Tab panel: give breathing room so content never clips under the tab bar */
[data-testid="stTabs"] [role="tabpanel"] {
    padding-top: 16px !important;
    overflow: visible !important;
}

/* ══ FIX 2: EXPANDER — flex layout so icon+text never overlap ══ */
[data-testid="stExpander"] summary {
    display: flex !important;
    flex-direction: row !important;
    align-items: center !important;
    gap: 10px !important;
    padding: 14px 18px !important;
    list-style: none !important;
    cursor: pointer !important;
    min-height: 50px !important;
}
[data-testid="stExpander"] summary::-webkit-details-marker { display: none !important; }
[data-testid="stExpander"] summary::marker { display: none !important; }
[data-testid="stExpander"] summary > * { flex-shrink: 0; }
[data-testid="stExpander"] summary p,
[data-testid="stExpander"] summary span {
    flex: 1 1 auto !important;
    min-width: 0 !important;
    text-transform: none !important;
    letter-spacing: normal !important;
    font-size: 14px !important;
    font-weight: 500 !important;
    white-space: normal !important;
    word-break: break-word !important;
}

/* ══ FIX 3: FILE UPLOADER — button text never overlaps ══ */
[data-testid="stFileUploadDropzone"] button,
[data-testid="stFileUploader"] button {
    display: inline-flex !important;
    align-items: center !important;
    justify-content: center !important;
    white-space: nowrap !important;
    overflow: visible !important;
    padding: 8px 20px !important;
    height: auto !important;
    min-width: fit-content !important;
    font-size: 14px !important;
}

/* ══ COMPONENT STYLES (custom HTML divs) ══ */
.hero {
    position: relative;
    background: linear-gradient(135deg, rgba(13,17,23,0.9) 0%, rgba(10,22,40,0.9) 100%);
    border: 1px solid rgba(0,255,170,0.2);
    border-radius: 20px; padding: 36px 44px; margin-bottom: 28px; overflow: visible;
}
.hero::after {
    content: ''; position: absolute; top: -1px; left: 40px; right: 40%; height: 1px;
    background: linear-gradient(90deg, transparent, rgba(0,255,170,0.6), transparent);
}
.hero-eyebrow { font-size: 10px; font-weight: 600; letter-spacing: .2em; text-transform: uppercase; color: rgba(0,255,170,0.7); margin-bottom: 12px; }
.hero h1 { font-size: 36px; font-weight: 800; margin: 0 0 10px; line-height: 1.15; }
.hero h1 span { color: #00ffaa; }
.hero p { margin: 0 0 18px; font-size: 14px; line-height: 1.6; }
.hero .badges { display: flex; gap: 8px; flex-wrap: wrap; }
.hero .badge { display: inline-flex; align-items: center; gap: 6px; background: rgba(0,255,170,0.08); border: 1px solid rgba(0,255,170,0.2); border-radius: 6px; padding: 5px 14px; font-size: 10px; font-weight: 700; letter-spacing: .1em; text-transform: uppercase; color: rgba(0,255,170,0.8); }

.card { background: rgba(255,255,255,0.03); border: 1px solid rgba(255,255,255,0.08); border-radius: 16px; padding: 24px 28px; margin-bottom: 16px; }
.stat-card { background: rgba(255,255,255,0.03); border: 1px solid rgba(255,255,255,0.07); border-radius: 14px; padding: 20px; text-align: center; }
.stat-card .val { font-size: 36px; font-weight: 800; color: #00ffaa; line-height: 1; }
.stat-card .lbl { font-size: 11px; opacity: 0.5; text-transform: uppercase; letter-spacing: .08em; margin-top: 6px; }

.arch-card { background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.06); border-radius: 14px; padding: 20px; text-align: center; }
.arch-card .arch-icon { font-size: 32px; margin-bottom: 10px; }
.arch-card .arch-title { font-size: 12px; font-weight: 700; letter-spacing: .12em; text-transform: uppercase; color: #00ffaa; margin-bottom: 8px; }
.arch-card .arch-body { font-size: 12px; opacity: 0.5; line-height: 1.55; }

.q-card { background: rgba(255,255,255,0.03); border: 1px solid rgba(0,255,170,0.2); border-radius: 18px; padding: 30px 34px; margin: 18px 0; position: relative; overflow: visible; }
.q-card::before { content: ''; position: absolute; top: 0; left: 0; width: 3px; height: 100%; background: linear-gradient(180deg, #00ffaa, transparent); border-radius: 18px 0 0 18px; }
.q-badge { display: inline-flex; align-items: center; gap: 5px; padding: 4px 14px; border-radius: 6px; font-size: 10px; font-weight: 700; letter-spacing: .1em; text-transform: uppercase; margin-bottom: 14px; }
.q-hr   { background: rgba(245,158,11,0.12); color: #fbbf24; border: 1px solid rgba(245,158,11,0.25); }
.q-tech { background: rgba(0,255,170,0.1); color: #00ffaa; border: 1px solid rgba(0,255,170,0.25); }
.q-num  { font-size: 11px; opacity: 0.4; margin-bottom: 8px; font-weight: 500; }
.q-text { font-size: 18px; font-weight: 600; line-height: 1.6; }
.q-hint { font-size: 12px; opacity: 0.5; margin-top: 14px; font-style: italic; padding: 10px 14px; background: rgba(255,255,255,0.03); border-radius: 8px; border-left: 2px solid rgba(0,255,170,0.2); }

.diff { padding: 4px 12px; border-radius: 6px; font-size: 11px; font-weight: 700; display: inline-block; letter-spacing: .05em; }
.d-easy   { background: rgba(0,255,170,0.1); color: #00ffaa; border: 1px solid rgba(0,255,170,0.2); }
.d-medium { background: rgba(245,158,11,0.1); color: #fbbf24; border: 1px solid rgba(245,158,11,0.2); }
.d-hard   { background: rgba(239,68,68,0.1); color: #f87171; border: 1px solid rgba(239,68,68,0.2); }

.pill-g { background: rgba(0,255,170,0.1); color: #00ffaa; padding: 5px 14px; border-radius: 20px; font-size: 12px; display: inline-block; margin: 3px; border: 1px solid rgba(0,255,170,0.2); }
.pill-r { background: rgba(239,68,68,0.1); color: #f87171; padding: 5px 14px; border-radius: 20px; font-size: 12px; display: inline-block; margin: 3px; border: 1px solid rgba(239,68,68,0.2); }
.pill-b { background: rgba(99,102,241,0.1); color: #a5b4fc; padding: 5px 14px; border-radius: 20px; font-size: 12px; display: inline-block; margin: 3px; border: 1px solid rgba(99,102,241,0.2); }

.status { background: rgba(0,255,170,0.05); border: 1px solid rgba(0,255,170,0.15); border-radius: 12px; padding: 12px 20px; font-size: 13px; font-weight: 500; margin-bottom: 20px; display: flex; align-items: center; gap: 12px; flex-wrap: wrap; row-gap: 6px; }
.status > span { white-space: nowrap; min-width: 0; overflow: hidden; text-overflow: ellipsis; }
.status .dot { width: 7px; height: 7px; border-radius: 50%; background: #00ffaa; display: inline-block; margin-right: 6px; box-shadow: 0 0 8px #00ffaa; animation: pulse-dot 2s infinite; }

.reasoning { background: rgba(0,255,170,0.04); border-left: 3px solid rgba(0,255,170,0.5); border-radius: 0 12px 12px 0; padding: 14px 18px; margin: 14px 0; font-size: 13px; color: rgba(0,255,170,0.8); }
.reasoning strong { color: #00ffaa; }

.cv-strength { background: rgba(0,255,170,0.05); border: 1px solid rgba(0,255,170,0.15); border-radius: 8px; padding: 10px 14px; margin: 5px 0; font-size: 13px; }
.cv-weakness { background: rgba(245,158,11,0.05); border: 1px solid rgba(245,158,11,0.15); border-radius: 8px; padding: 10px 14px; margin: 5px 0; font-size: 13px; }
.cv-gap      { background: rgba(239,68,68,0.05); border: 1px solid rgba(239,68,68,0.15); border-radius: 8px; padding: 10px 14px; margin: 5px 0; font-size: 13px; }

.msg-user { background: rgba(99,102,241,0.1); border: 1px solid rgba(99,102,241,0.2); border-radius: 16px 16px 4px 16px; padding: 14px 18px; margin: 10px 0; font-size: 14px; }
.msg-ai   { background: rgba(0,255,170,0.05); border: 1px solid rgba(0,255,170,0.12); border-radius: 16px 16px 16px 4px; padding: 14px 18px; margin: 10px 0; font-size: 14px; }

.score-display { background: linear-gradient(135deg, rgba(0,255,170,0.1), rgba(0,200,150,0.05)); border: 1px solid rgba(0,255,170,0.25); border-radius: 18px; padding: 28px; text-align: center; }
.score-display .big-num { font-size: 56px; font-weight: 800; color: #00ffaa; line-height: 1; }
.score-display .rating { font-size: 16px; font-weight: 600; margin-top: 8px; }
</style>
"""


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 · SESSION STATE INIT
# ══════════════════════════════════════════════════════════════════════════════
def init():
    defaults = {
        "memory":           None,
        "brain":            None,
        "groq_client":      None,
        "q_pool":           [],
        "current_q":        None,
        "session_active":   False,
        "phase":            "setup",
        "cv_text":          "",
        "cv_analysis":      None,
        "cv_rewrite":       None,
        "company_desc":     "",
        "q_count":          20,
        "show_eval":        False,
        "last_eval":        None,
        "last_scores":      None,
        "last_reasoning":   [],
        "perception":       None,
        "decision":         None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    if st.session_state.memory is None:
        st.session_state.memory = AgentMemory()

    # Always rebuild client so a freshly-entered key is picked up immediately
    current_key = get_api_key()
    if not current_key:
        get_groq_client()   # triggers the key-entry screen + st.stop()

    # Rebuild client & brain whenever the key changes or they were cleared
    cached_client = st.session_state.get("groq_client")
    key_changed = (
        cached_client is None or
        getattr(cached_client, "_api_key", None) != current_key
    )
    if key_changed:
        new_client = Groq(api_key=current_key)
        new_client._api_key = current_key          # tag for change-detection
        st.session_state.groq_client = new_client
        st.session_state.brain = None              # force brain rebuild

    if st.session_state.brain is None:
        st.session_state.brain = AgentBrain(
            st.session_state.groq_client, st.session_state.memory
        )


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 · SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
def render_sidebar():
    mem: AgentMemory = st.session_state.memory
    cs = mem.current_session

    with st.sidebar:
        st.markdown("""
        <div style="padding:24px 16px 12px">
          <div style="display:flex;align-items:center;gap:12px;margin-bottom:4px">
            <div style="width:38px;height:38px;background:rgba(0,255,170,0.1);border:1px solid rgba(0,255,170,0.25);border-radius:10px;display:flex;align-items:center;justify-content:center;font-size:20px">🎯</div>
            <div>
              <div style="font-family:'Syne',sans-serif;font-size:16px;font-weight:800;color:#f0f6ff;letter-spacing:-.01em">Interview Agent</div>
              <div style="font-size:10px;color:rgba(0,255,170,0.6);font-weight:600;letter-spacing:.1em;text-transform:uppercase">Groq · LLaMA 3.3 70B</div>
            </div>
          </div>
        </div>
        <div style="height:1px;background:rgba(0,255,170,0.1);margin:0 0 12px"></div>
        """, unsafe_allow_html=True)

        if cs:
            answered = len(cs.get("answers", []))
            total = len(st.session_state.q_pool)
            pct = int(answered / total * 100) if total else 0
            diff = cs.get("difficulty", "medium")
            st.markdown(f"""
            <div style="background:rgba(0,255,170,0.06);border:1px solid rgba(0,255,170,0.18);border-radius:12px;padding:14px;margin-bottom:14px">
              <div style="color:rgba(0,255,170,0.6);font-size:10px;text-transform:uppercase;letter-spacing:.1em;margin-bottom:8px;display:flex;align-items:center;gap:6px"><span style="width:6px;height:6px;background:#00ffaa;border-radius:50%;display:inline-block;box-shadow:0 0 6px #00ffaa"></span> Active Session</div>
              <div style="font-size:13px;color:#f0f6ff;font-weight:600">{cs.get('role','')} <span style="color:rgba(200,214,229,0.45);font-weight:400">@ {cs.get('company','')}</span></div>
              <div style="color:rgba(200,214,229,0.4);font-size:11px;margin:6px 0">{answered}/{total} questions answered</div>
              <div style="background:rgba(255,255,255,0.06);border-radius:6px;height:4px;margin-bottom:8px">
                <div style="width:{pct}%;background:linear-gradient(90deg,#00ffaa,#00c896);height:100%;border-radius:6px;transition:width .6s ease"></div>
              </div>
              <span class="diff d-{diff}" style="font-size:10px">⚡ {diff.capitalize()}</span>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style="background:rgba(255,255,255,0.03);border:1px solid rgba(255,255,255,0.07);border-radius:12px;padding:12px;margin-bottom:14px">
              <div style="color:rgba(200,214,229,0.3);font-size:10px;text-transform:uppercase;letter-spacing:.1em;margin-bottom:6px">Agent</div>
              <div style="color:rgba(200,214,229,0.3);font-size:13px">⏸ No active session</div>
            </div>
            """, unsafe_allow_html=True)

        n_sessions = len(mem.sessions)
        n_weak     = len(mem.weak_skills)
        n_asked    = len(mem._data.get("question_history",[]))
        n_cv_gaps  = len(mem.cv_gaps)

        st.markdown(f"""
        <div style="background:rgba(255,255,255,0.03);border:1px solid rgba(255,255,255,0.07);border-radius:12px;padding:14px;margin-bottom:14px">
          <div style="color:rgba(0,255,170,0.6);font-size:10px;text-transform:uppercase;letter-spacing:.1em;margin-bottom:10px">🧠 Agent Memory</div>
          <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px">
            <div style="text-align:center;background:rgba(0,255,170,0.05);border-radius:8px;padding:10px">
              <div style="color:#00ffaa;font-size:22px;font-weight:800;font-family:'Syne',sans-serif">{n_sessions}</div>
              <div style="color:rgba(200,214,229,0.4);font-size:10px;text-transform:uppercase;letter-spacing:.06em">Sessions</div>
            </div>
            <div style="text-align:center;background:rgba(245,158,11,0.05);border-radius:8px;padding:10px">
              <div style="color:#fbbf24;font-size:22px;font-weight:800;font-family:'Syne',sans-serif">{n_weak}</div>
              <div style="color:rgba(200,214,229,0.4);font-size:10px;text-transform:uppercase;letter-spacing:.06em">Weak Skills</div>
            </div>
            <div style="text-align:center;background:rgba(99,102,241,0.05);border-radius:8px;padding:10px">
              <div style="color:#a5b4fc;font-size:22px;font-weight:800;font-family:'Syne',sans-serif">{n_asked}</div>
              <div style="color:rgba(200,214,229,0.4);font-size:10px;text-transform:uppercase;letter-spacing:.06em">Qs Seen</div>
            </div>
            <div style="text-align:center;background:rgba(239,68,68,0.05);border-radius:8px;padding:10px">
              <div style="color:#f87171;font-size:22px;font-weight:800;font-family:'Syne',sans-serif">{n_cv_gaps}</div>
              <div style="color:rgba(200,214,229,0.4);font-size:10px;text-transform:uppercase;letter-spacing:.06em">CV Gaps</div>
            </div>
          </div>
        </div>
        """, unsafe_allow_html=True)

        if mem.weak_skills:
            top3 = sorted(mem.weak_skills.items(), key=lambda x: x[1], reverse=True)[:4]
            max_v = max(v for _,v in top3) or 1
            st.markdown('<div style="color:rgba(0,255,170,0.6);font-size:10px;text-transform:uppercase;letter-spacing:.1em;margin-bottom:8px">🎯 Top Weak Areas</div>', unsafe_allow_html=True)
            for tag, val in top3:
                bar = int(val / max_v * 100)
                st.markdown(f"""
                <div style="margin-bottom:8px">
                  <div style="display:flex;justify-content:space-between;align-items:center;font-size:11px;color:rgba(200,214,229,0.5);margin-bottom:3px">
                    <span style="overflow:hidden;text-overflow:ellipsis;white-space:nowrap;min-width:0;flex:1;margin-right:6px">{tag.replace('_',' ').title()[:20]}</span>
                    <span style="color:#fbbf24;flex-shrink:0">{val:.1f}</span>
                  </div>
                  <div style="background:rgba(255,255,255,0.06);border-radius:4px;height:3px">
                    <div style="width:{bar}%;background:linear-gradient(90deg,#fbbf24,#f59e0b);height:100%;border-radius:4px"></div>
                  </div>
                </div>""", unsafe_allow_html=True)
            st.markdown("<div style='height:1px;background:rgba(0,255,170,0.1);margin:10px 0'></div>", unsafe_allow_html=True)

        st.markdown('<div style="color:rgba(0,255,170,0.6);font-size:10px;text-transform:uppercase;letter-spacing:.1em;margin-bottom:8px">⚙️ Settings</div>', unsafe_allow_html=True)
        st.session_state.q_count = st.select_slider(
            "Questions per session",
            options=[10, 20, 30, 50, 100, 200],
            value=st.session_state.q_count,
        )

        st.markdown("<div style='height:1px;background:rgba(0,255,170,0.1);margin:10px 0'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 New Session", use_container_width=True):
                mem.close_session()
                st.session_state.phase = "setup"
                st.session_state.session_active = False
                st.session_state.q_pool = []
                st.session_state.current_q = None
                st.session_state.show_eval = False
                st.session_state.cv_analysis = None
                st.session_state.cv_rewrite = None
                st.rerun()
        with col2:
            if st.button("🗑️ Reset All", use_container_width=True):
                mem.reset()
                for k in ["session_active","q_pool","current_q","phase","cv_text",
                          "cv_analysis","cv_rewrite","show_eval","last_eval","last_scores"]:
                    st.session_state[k] = None if k not in ["q_pool"] else []
                st.session_state.phase = "setup"
                st.session_state.session_active = False
                st.rerun()

        # API Key management — always visible
        st.markdown("<div style='height:1px;background:rgba(0,255,170,0.1);margin:10px 0'></div>", unsafe_allow_html=True)
        api_set = bool(get_api_key())
        if api_set:
            st.markdown("""
            <div style="background:rgba(0,255,170,0.06);border:1px solid rgba(0,255,170,0.2);border-radius:10px;padding:10px 14px;font-size:12px;color:rgba(0,255,170,0.8)">
              🔑 API Key: Connected
            </div>""", unsafe_allow_html=True)
            if st.button("🔑 Change API Key", use_container_width=True):
                st.session_state["_groq_api_key"] = ""
                st.session_state["groq_client"] = None
                st.session_state["brain"] = None
                st.rerun()
        else:
            st.markdown("""
            <div style="background:rgba(239,68,68,0.06);border:1px solid rgba(239,68,68,0.2);border-radius:10px;padding:10px 14px;font-size:12px;color:#f87171">
              ⚠️ No API Key set
            </div>""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 · TAB: JOB SETUP
# ══════════════════════════════════════════════════════════════════════════════
def tab_setup():
    mem: AgentMemory   = st.session_state.memory
    brain: AgentBrain  = st.session_state.brain

    st.markdown("""
    <div class="hero">
      <div class="hero-eyebrow">AI-Powered Interview System</div>
      <h1>Your Personal <span>Interview Coach</span></h1>
      <p>An autonomous AI agent that perceives, reasons, acts, and learns — adapting to your strengths and weaknesses in real-time. Powered by Groq LLaMA 3.3 70B.</p>
      <div class="badges">
        <span class="badge">Perceive</span>
        <span class="badge">Reason</span>
        <span class="badge">Act</span>
        <span class="badge">Learn</span>
      </div>
    </div>
    """, unsafe_allow_html=True)

    with st.expander("🤖 Agent Architecture — How This Works", expanded=False):
        c1, c2, c3, c4 = st.columns(4)
        for col, icon, title, body in [
            (c1,"👁️","PERCEIVE","Reads job description, CV, past scores, and weak skills to build a world model."),
            (c2,"🧠","REASON", "Applies decision rules: if tech score < 3 → more technical questions; if CV gap → target that skill."),
            (c3,"⚡","ACT",    "Generates tailored questions, evaluates answers, rewrites CV, and produces reports."),
            (c4,"📚","LEARN",  "After every answer, updates skill deficit scores and adapts future behavior permanently."),
        ]:
            col.markdown(f"""
            <div class="arch-card">
              <div class="arch-icon">{icon}</div>
              <div class="arch-title">{title}</div>
              <div class="arch-body">{body}</div>
            </div>""", unsafe_allow_html=True)

    st.markdown("### 📋 Job Details")
    col1, col2 = st.columns([3, 2])
    with col1:
        company = st.text_input("🏢 Company Name", placeholder="e.g. Google, HSBC, Tesla",
                                key="inp_company")
        role    = st.text_input("💼 Job Role / Title",
                                placeholder="e.g. Senior Data Scientist", key="inp_role")
    with col2:
        level = st.selectbox("📊 Position Level",
                             ["Intern","Junior","Mid-Level","Senior",
                              "Lead / Principal","Manager / Director"],
                             key="inp_level")
        st.markdown(f"""
        <div class="stat-card" style="margin-top:12px">
          <div class="val">{st.session_state.q_count}</div>
          <div class="lbl">Questions (change in sidebar)</div>
        </div>""", unsafe_allow_html=True)

    responsibilities = st.text_area(
        "📌 Key Responsibilities",
        placeholder="Describe what this role does day-to-day...",
        height=100, key="inp_resp"
    )
    requirements = st.text_area(
        "✅ Required Skills / Qualifications",
        placeholder="e.g. 5+ years Python, experience with MLOps, SQL...",
        height=100, key="inp_reqs"
    )

    c_btn, c_clr, _ = st.columns([2, 1.5, 4])
    with c_btn:
        go_btn = st.button(
            "🚀 Start Interview Session", type="primary",
            use_container_width=True,
            disabled=not (company and role and responsibilities and requirements)
        )
    with c_clr:
        if st.button("🗑️ Clear", use_container_width=True):
            for k in ["inp_company","inp_role","inp_resp","inp_reqs"]:
                st.session_state.pop(k, None)
            st.rerun()

    if go_btn:
        with st.spinner("🧠 Agent perceiving inputs and reasoning about your session..."):
            desc = brain.act_research_company(company, role)
            st.session_state.company_desc = desc

            mem.start_session(company, role, level, responsibilities, requirements)

            perception = brain.perceive()
            st.session_state.perception = perception

            decision = brain.reason(perception)
            if mem.score_history:
                decision["difficulty"] = _recommend_difficulty(mem)
            st.session_state.decision = decision
            mem.update_session("difficulty", decision["difficulty"])
            st.session_state.last_reasoning = decision.get("rationale", [])

            st.info(f"📝 Generating {st.session_state.q_count} questions "
                    f"(difficulty: {decision['difficulty']})…")
            q_data = brain.act_generate_questions(
                perception, decision, count=st.session_state.q_count
            )

        hr_qs   = q_data.get("hr_questions", [])
        tech_qs = q_data.get("technical_questions", [])
        all_qs  = hr_qs + tech_qs

        if not all_qs:
            st.error("No questions generated. Please check your GROQ_API_KEY.")
            return

        mem.update_session("question_pool", all_qs)
        st.session_state.q_pool = all_qs
        st.session_state.session_active = True
        st.session_state.phase = "interview"

        first_q = brain.select_next_question(all_qs)
        st.session_state.current_q = first_q

        rationale_html = " · ".join(decision.get("rationale", []))
        st.success(f"✅ {len(all_qs)} questions ready · Difficulty: **{decision['difficulty'].capitalize()}**")
        st.markdown(f'<div class="reasoning"><strong>🧠 Agent Reasoning:</strong> {rationale_html}</div>',
                    unsafe_allow_html=True)
        import time; time.sleep(1.5)
        st.rerun()


def _recommend_difficulty(mem: AgentMemory) -> str:
    recent = mem.score_history[-3:]
    if not recent: return "medium"
    avg = sum(s.get("overall", 3) for s in recent) / len(recent)
    if avg >= 4.0: return "hard"
    if avg <= 2.5: return "easy"
    return "medium"


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 · TAB: CV ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
def tab_cv():
    mem: AgentMemory  = st.session_state.memory
    brain: AgentBrain = st.session_state.brain
    cs = mem.current_session

    st.markdown("## 📄 CV Analysis & Optimization")

    if not cs:
        st.warning("⚠️ Please start a session in **Job Setup** first so the agent knows the target role.")

    col_up, col_job = st.columns([2, 3])

    with col_up:
        st.markdown("### 📎 Upload Your CV")
        uploaded = st.file_uploader("PDF or DOCX", type=["pdf","docx"])
        if uploaded:
            with st.spinner("Extracting text..."):
                text = extract_cv_text(uploaded)
            if text:
                st.session_state.cv_text = text
                st.success(f"✅ Extracted {len(text.split())} words")
                with st.expander("Preview Extracted Text"):
                    st.text_area("CV Text", text[:2000], height=200)
            else:
                st.error("Could not extract text from this file.")

    with col_job:
        st.markdown("### 🎯 Job Context")
        if cs:
            st.markdown(f"""
            <div class="card">
              <div style="font-size:13px;color:rgba(200,214,229,0.5);margin-bottom:8px">Current Session</div>
              <div style="font-size:16px;font-weight:700">{cs.get('role','')} <span style="color:rgba(200,214,229,0.5);font-weight:400">@ {cs.get('company','')}</span></div>
              <div style="font-size:13px;color:rgba(200,214,229,0.5);margin-top:4px">{cs.get('level','')}</div>
            </div>
            """, unsafe_allow_html=True)
            if st.session_state.company_desc:
                with st.expander("🏢 Company Research"):
                    st.markdown(st.session_state.company_desc)
        else:
            role_m  = st.text_input("Role (manual)", placeholder="e.g. Data Scientist")
            level_m = st.selectbox("Level", ["Intern","Junior","Mid-Level","Senior","Lead"])
            resp_m  = st.text_area("Responsibilities", height=80)
            reqs_m  = st.text_area("Requirements", height=80)
            if st.button("Save Job Context"):
                mem.start_session("Unknown Company", role_m, level_m, resp_m, reqs_m)
                st.rerun()

    st.divider()

    c1, c2, _ = st.columns([2, 2, 4])
    with c1:
        analyse_btn = st.button(
            "🔍 Analyse CV vs Job", type="primary", use_container_width=True,
            disabled=not (st.session_state.cv_text and cs)
        )
    with c2:
        rewrite_btn = False
        if st.session_state.cv_analysis:
            rewrite_btn = st.button("✍️ Rewrite CV", use_container_width=True)

    if analyse_btn:
        cs = mem.current_session
        with st.spinner("🧠 Agent analysing your CV against job requirements..."):
            analysis = brain.act_analyse_cv(
                st.session_state.cv_text,
                cs.get("role",""), cs.get("level",""),
                cs.get("responsibilities",""), cs.get("requirements","")
            )
        st.session_state.cv_analysis = analysis
        gaps      = analysis.get("missing_skills",[]) + analysis.get("experience_gaps",[])
        strengths = analysis.get("strengths",[])
        mem.store_cv_analysis(gaps, strengths)
        st.rerun()

    if rewrite_btn:
        cs = mem.current_session
        with st.spinner("✍️ Rewriting your CV to match the job..."):
            rewrite = brain.act_rewrite_cv(
                st.session_state.cv_text,
                cs.get("role",""), cs.get("level",""),
                cs.get("responsibilities",""), cs.get("requirements",""),
                st.session_state.cv_analysis,
            )
        st.session_state.cv_rewrite = rewrite
        st.rerun()

    if st.session_state.cv_analysis:
        an = st.session_state.cv_analysis
        st.markdown("### 📊 Analysis Results")

        col_gauge, col_breakdown = st.columns([2, 3])
        with col_gauge:
            st.plotly_chart(cv_match_gauge(an.get("relevance_score",0)),
                            use_container_width=True)
        with col_breakdown:
            mb = an.get("match_breakdown", {})
            for label, key, colour in [
                ("Technical Match",  "technical_match",  "#00ffaa"),
                ("Experience Match", "experience_match", "#a5b4fc"),
                ("Keyword Match",    "keyword_match",    "#fbbf24"),
            ]:
                v = mb.get(key, 0) * 10
                st.markdown(f"""
                <div style="margin-bottom:14px">
                  <div style="display:flex;justify-content:space-between;margin-bottom:5px">
                    <span style="font-size:13px;font-weight:500;color:rgba(200,214,229,0.7)">{label}</span>
                    <span style="font-size:13px;font-weight:700;color:{colour}">{mb.get(key,0)}/10</span>
                  </div>
                  <div style="background:rgba(255,255,255,0.06);border-radius:999px;height:6px">
                    <div style="width:{v}%;background:{colour};height:100%;border-radius:999px"></div>
                  </div>
                </div>""", unsafe_allow_html=True)

        c_str, c_wk, c_gap = st.columns(3)
        with c_str:
            st.markdown("**✅ Strengths**")
            for s in an.get("strengths",[]):
                st.markdown(f'<div class="cv-strength">✓ {s}</div>', unsafe_allow_html=True)
        with c_wk:
            st.markdown("**⚠️ Weaknesses**")
            for w in an.get("weaknesses",[]):
                st.markdown(f'<div class="cv-weakness">⚠ {w}</div>', unsafe_allow_html=True)
        with c_gap:
            st.markdown("**🔴 Missing Skills**")
            for g in an.get("missing_skills",[]):
                st.markdown(f'<div class="cv-gap">✗ {g}</div>', unsafe_allow_html=True)

        col_kw, col_ats = st.columns(2)
        with col_kw:
            st.markdown("**🔑 Missing Keywords (for ATS)**")
            kws = an.get("keywords_missing",[])
            if kws:
                kw_html = " ".join(f'<span class="pill-r">{k}</span>' for k in kws)
                st.markdown(kw_html, unsafe_allow_html=True)
        with col_ats:
            st.markdown("**🤖 ATS Issues**")
            for issue in an.get("ats_issues",[]):
                st.markdown(f"• {issue}")

        st.markdown("**💡 Recommendations**")
        for rec in an.get("recommendations",[]):
            st.markdown(f'<span class="pill-b">→ {rec}</span>', unsafe_allow_html=True)

        if an.get("overall_assessment"):
            st.info(f"📌 {an['overall_assessment']}")

        st.markdown(f"""
        <div class="reasoning">
          <strong>🧠 Agent Memory Updated:</strong>
          {len(an.get('missing_skills',[]))} CV gaps added to weak skills.
          Future questions will target: {', '.join(an.get('missing_skills',[])[:3])}.
        </div>
        """, unsafe_allow_html=True)

    if st.session_state.cv_rewrite:
        st.divider()
        st.markdown("### ✍️ Rewritten CV")
        st.markdown("""
        <div style="background:rgba(245,158,11,0.08);border:1px solid rgba(245,158,11,0.3);border-radius:10px;padding:12px 16px;margin-bottom:12px;font-size:13px;color:#fbbf24">
          ⚠️ The agent has improved wording, structure, and keywords. No experience has been fabricated.
        </div>""", unsafe_allow_html=True)

        with st.expander("📄 View Rewritten CV", expanded=True):
            st.text_area("Rewritten CV", st.session_state.cv_rewrite, height=500)

        c_pdf, c_docx, _ = st.columns([2, 2, 4])
        with c_pdf:
            pdf_bytes = generate_cv_pdf(st.session_state.cv_rewrite)
            st.download_button(
                "⬇️ Download as PDF", data=pdf_bytes,
                file_name="rewritten_cv.pdf", mime="application/pdf",
                use_container_width=True
            )
        with c_docx:
            doc_obj = DocxDocument()
            for line in st.session_state.cv_rewrite.split("\n"):
                doc_obj.add_paragraph(line)
            docx_buf = io.BytesIO()
            doc_obj.save(docx_buf)
            st.download_button(
                "⬇️ Download as DOCX", data=docx_buf.getvalue(),
                file_name="rewritten_cv.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 · TAB: INTERVIEW
# ══════════════════════════════════════════════════════════════════════════════
def tab_interview():
    mem: AgentMemory  = st.session_state.memory
    brain: AgentBrain = st.session_state.brain
    cs = mem.current_session

    if not cs or st.session_state.phase != "interview":
        st.info("👈 Set up your job in the **Job Setup** tab to begin the interview.")
        return

    answered = len(cs.get("answers", []))
    total    = len(st.session_state.q_pool)
    diff     = cs.get("difficulty", "medium")

    scores_so_far = cs.get("scores", [])
    if scores_so_far:
        avg = sum((s["technical_knowledge"]+s["communication"]+s["confidence"])/3
                  for s in scores_so_far) / len(scores_so_far)
        avg_str = f"Avg: {avg:.1f}/5"
    else:
        avg_str = ""

    st.markdown(f"""
    <div class="status">
      <span><span class="dot"></span>Agent Active</span>
      <span>🏢 {cs.get('company','')} · {cs.get('role','')}</span>
      <span>📋 {answered}/{total}</span>
      <span class="diff d-{diff}">⚡ {diff.capitalize()}</span>
      {f'<span>📊 {avg_str}</span>' if avg_str else ''}
    </div>
    """, unsafe_allow_html=True)

    st.progress(answered / total if total else 0,
                text=f"Progress: {answered}/{total} questions answered")

    if st.session_state.last_reasoning:
        rationale = " · ".join(st.session_state.last_reasoning)
        st.markdown(f'<div class="reasoning"><strong>🧠 Agent Decision Logic:</strong> {rationale}</div>',
                    unsafe_allow_html=True)

    if st.session_state.show_eval and st.session_state.last_eval:
        ev = st.session_state.last_eval
        sc = st.session_state.last_scores

        st.markdown("---")
        st.markdown("### 📊 Evaluation of Your Last Answer")

        col_bars, col_radar = st.columns([3, 2])
        with col_bars:
            score_bars(sc)
        with col_radar:
            overall = (sc["technical_knowledge"]+sc["communication"]+sc["confidence"])/3
            st.markdown(f"""
            <div class="score-display">
              <div class="big-num">{overall:.1f}</div>
              <div style="font-size:11px;color:rgba(200,214,229,0.4);text-transform:uppercase;letter-spacing:.1em;margin:6px 0">Overall Score</div>
              <div class="rating">
                {"🌟 Excellent" if overall>=4.5 else "👍 Good" if overall>=3.5 else "📈 Average" if overall>=2.5 else "📚 Needs Work"}
              </div>
            </div>""", unsafe_allow_html=True)

        c_str, c_wk = st.columns(2)
        with c_str:
            st.markdown("**✅ Strengths**")
            for s in ev.get("strengths",[]): st.markdown(f'<span class="pill-g">✓ {s}</span>', unsafe_allow_html=True)
        with c_wk:
            st.markdown("**⚠️ Weaknesses**")
            for w in ev.get("weaknesses",[]): st.markdown(f'<span class="pill-r">✗ {w}</span>', unsafe_allow_html=True)

        st.markdown("**💡 Suggestions**")
        for sg in ev.get("suggestions",[]): st.markdown(f'<span class="pill-b">→ {sg}</span>', unsafe_allow_html=True)

        if ev.get("ideal_answer_summary"):
            with st.expander("📌 Ideal Answer Guide"):
                st.info(ev["ideal_answer_summary"])
        if ev.get("overall_comment"):
            st.markdown(f"> _{ev['overall_comment']}_")
        st.markdown("---")

    cq = st.session_state.current_q

    if cq is None or answered >= total:
        st.balloons()
        st.success("🎉 Interview complete! Check your full report in the **Dashboard** tab.")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📊 View Dashboard", type="primary", use_container_width=True):
                mem.close_session()
                st.session_state.phase = "done"
                st.rerun()
        with col2:
            if st.button("🔄 New Session", use_container_width=True):
                mem.close_session()
                st.session_state.phase = "setup"
                st.session_state.session_active = False
                st.session_state.q_pool = []
                st.session_state.current_q = None
                st.session_state.show_eval = False
                st.rerun()
        return

    q_type = cq.get("type","technical")
    badge_cls = "q-hr" if q_type == "hr" else "q-tech"
    type_label = "Behavioural / HR" if q_type == "hr" else "Technical"
    tags_str = " · ".join(cq.get("tags",[]))

    st.markdown(f"""
    <div class="q-card">
      <div class="q-num">Question {answered+1} of {total}</div>
      <span class="q-badge {badge_cls}">{type_label}</span>
      {"&nbsp;<span style='font-size:11px;color:rgba(200,214,229,0.5)'>" + tags_str + "</span>" if tags_str else ""}
      <div class="q-text">{cq.get('question','')}</div>
      {"<div class='q-hint'>💡 Consider: " + cq.get('follow_up_hint','') + "</div>" if cq.get('follow_up_hint') else ""}
    </div>
    """, unsafe_allow_html=True)

    if q_type == "hr" and st.session_state.company_desc:
        with st.expander("🏢 Company Context"):
            st.markdown(st.session_state.company_desc)

    answer = st.text_area(
        "✏️ Your Answer",
        placeholder="Type your full answer here. Be specific and structured.",
        height=200, key=f"ans_{answered}",
    )

    col1, col2, col3, col4 = st.columns([2, 1.5, 1.5, 1.5])
    with col1:
        submit = st.button("📤 Submit Answer", type="primary",
                           use_container_width=True, disabled=not answer.strip())
    with col2:
        skip = st.button("⏭️ Skip", use_container_width=True)
    with col3:
        regen = st.button("🔁 Next Question", use_container_width=True)
    with col4:
        if st.button("🛑 End Session", use_container_width=True):
            mem.close_session()
            st.session_state.phase = "done"
            st.rerun()

    if submit and answer.strip():
        with st.spinner("🧠 Agent evaluating..."):
            ev = brain.act_evaluate_answer(
                cq, answer,
                cs.get("role",""), cs.get("level",""),
                context=st.session_state.company_desc,
            )
        sc = ev.get("scores", {"technical_knowledge":3,"communication":3,"confidence":3})

        mem.mark_asked(cq["question"])
        mem.record_qa(cq, answer, ev, sc)
        brain.learn(cq, sc)

        perception = brain.perceive()
        decision   = brain.reason(perception)
        st.session_state.last_reasoning = decision.get("rationale", [])
        mem.update_session("difficulty", decision["difficulty"])

        st.session_state.last_eval   = ev
        st.session_state.last_scores = sc
        st.session_state.show_eval   = True

        remaining = [q for q in st.session_state.q_pool
                     if q.get("question") != cq.get("question")]
        next_q = brain.select_next_question(remaining)
        st.session_state.current_q = next_q

        st.rerun()

    if skip or regen:
        remaining = [q for q in st.session_state.q_pool
                     if q.get("question") != cq.get("question")]
        next_q = brain.select_next_question(remaining)
        st.session_state.current_q = next_q
        st.session_state.show_eval = False
        st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 · TAB: DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
def tab_dashboard():
    mem: AgentMemory = st.session_state.memory

    st.markdown("## 📊 Performance Dashboard")

    all_sessions = mem.sessions
    cs_data = all_sessions[-1] if all_sessions else mem.current_session
    score_hist = mem.score_history

    if not cs_data:
        st.info("Complete your first interview session to see data here.")
        return

    summary = cs_data.get("summary") or AgentMemory._summarise_scores(cs_data.get("scores", []))

    if summary:
        k1, k2, k3, k4 = st.columns(4)
        for col, label, key, icon in [
            (k1, "Technical",      "technical_knowledge", "🔧"),
            (k2, "Communication",  "communication",       "🗣️"),
            (k3, "Confidence",     "confidence",          "💪"),
            (k4, "Overall",        "overall",             "🏆"),
        ]:
            val = summary.get(key, 0)
            col.metric(f"{icon} {label}", f"{val:.2f}/5", delta=_rating_str(val))

    st.divider()

    col_r, col_w = st.columns(2)
    with col_r:
        if summary:
            st.plotly_chart(radar_chart(summary, "Session Performance"),
                            use_container_width=True)
    with col_w:
        wf = weakness_chart(mem.weak_skills)
        if wf:
            st.plotly_chart(wf, use_container_width=True)
        elif mem.cv_gaps:
            st.markdown("**CV Gaps (to target)**")
            for g in mem.cv_gaps: st.markdown(f"• {g}")

    if len(score_hist) > 1:
        tf = trend_chart(score_hist)
        if tf:
            st.plotly_chart(tf, use_container_width=True)

    if score_hist:
        st.markdown("### 📋 Session History")
        df = pd.DataFrame([{
            "#": i+1, "Date": s.get("timestamp","")[:10],
            "Company": s.get("company",""), "Role": s.get("role",""),
            "Technical": f"{s.get('technical_knowledge',0):.2f}",
            "Communication": f"{s.get('communication',0):.2f}",
            "Confidence": f"{s.get('confidence',0):.2f}",
            "Overall": f"{s.get('overall',0):.2f}",
        } for i, s in enumerate(score_hist)])
        st.dataframe(df, use_container_width=True, hide_index=True)

    st.markdown("### 📝 Last Session Q&A Review")
    questions = cs_data.get("questions",[])
    answers   = cs_data.get("answers",[])
    evals     = cs_data.get("evaluations",[])
    scores    = cs_data.get("scores",[])

    for i, (q, a, ev, sc) in enumerate(zip(questions, answers, evals, scores)):
        overall_q = (sc.get("technical_knowledge",0) + sc.get("communication",0) + sc.get("confidence",0)) / 3
        with st.expander(f"Q{i+1} [{q.get('type','').upper()}]: {q.get('question','')[:70]}… — **{overall_q:.1f}/5**"):
            score_bars(sc)
            st.markdown(f"**Your Answer:** {a}")
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**Strengths**")
                for s in ev.get("strengths",[]): st.markdown(f"✓ {s}")
            with c2:
                st.markdown("**Weaknesses**")
                for w in ev.get("weaknesses",[]): st.markdown(f"✗ {w}")
            for sg in ev.get("suggestions",[]): st.markdown(f'<span class="pill-b">→ {sg}</span>', unsafe_allow_html=True)

    st.divider()
    c_exp, _ = st.columns([2, 5])
    with c_exp:
        if st.button("📄 Generate PDF Report", type="primary", use_container_width=True):
            with st.spinner("Generating PDF…"):
                pdf = generate_pdf_report(mem, cs_data)
            fname = f"interview_report_{cs_data.get('company','').replace(' ','_')}_{cs_data.get('timestamp','')[:10]}.pdf"
            st.download_button("⬇️ Download PDF", data=pdf, file_name=fname,
                               mime="application/pdf", use_container_width=True)


def _rating_str(v: float) -> str:
    if v >= 4.5: return "Excellent"
    if v >= 3.5: return "Good"
    if v >= 2.5: return "Average"
    return "Needs Work"


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 · TAB: CHATBOT
# ══════════════════════════════════════════════════════════════════════════════
def tab_chatbot():
    mem: AgentMemory  = st.session_state.memory
    brain: AgentBrain = st.session_state.brain
    cs = mem.current_session

    st.markdown("## 💬 Interview Coach Chatbot")
    st.markdown("Ask anything — interview tips, salary negotiation, career advice, CV help.")

    history = mem.chat_history
    for msg in history:
        cls = "msg-user" if msg["role"] == "user" else "msg-ai"
        prefix = "👤" if msg["role"] == "user" else "🤖"
        st.markdown(f'<div class="{cls}">{prefix} {msg["content"]}</div>',
                    unsafe_allow_html=True)

    user_msg = st.chat_input("Ask the AI Interview Coach…")
    if user_msg:
        context = f"Preparing for {cs.get('role','')} at {cs.get('company','')}" if cs else ""
        with st.spinner("Thinking…"):
            reply = brain.act_chat(user_msg, history, context)
        mem.add_chat("user", user_msg)
        mem.add_chat("assistant", reply)
        st.rerun()

    if history:
        if st.button("🗑️ Clear Chat"):
            mem.clear_chat()
            st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 14 · MAIN ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════
def main():
    init()
    render_sidebar()

    tabs = st.tabs([
        "🏠 Job Setup",
        "📄 CV Analysis",
        "🎤 Interview",
        "📊 Dashboard",
        "💬 Chatbot",
    ])
    with tabs[0]: tab_setup()
    with tabs[1]: tab_cv()
    with tabs[2]: tab_interview()
    with tabs[3]: tab_dashboard()
    with tabs[4]: tab_chatbot()


if __name__ == "__main__":
    main()
